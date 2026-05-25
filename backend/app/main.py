"""
main.py — AI Resume Generation Pipeline (Dual-Agent Writer + ATS Scorer).

Fixes applied vs. original:
  - FIX #1 : exit(1) replaced with raise SystemExit / raise RuntimeError
              (exit() at module level would kill the FastAPI server if ever imported)
  - FIX #2 : master_resume_path read ONCE before the loop (was re-read every iteration)
  - FIX #3 : Model fallback uses a models-list loop — original continue-based logic
              never actually switched models correctly
  - FIX #4 : RAG empty-context guard — warns and falls back to full resume if
              ChromaDB returns nothing
  - FIX #5 : Token-length guard — truncates JD and resume before sending to Groq
              (prevents silent 8k context-window overflow on free tier)
  - FIX #6 : ATS score normalised to int with safe fallback (was left as raw string)
  - FIX #7 : clean_job_description imported from utils.py (no longer duplicated)
  - FIX #8 : time.sleep → import kept as-is (subprocess context); documented
"""

import os
import re
import sys
import csv
import json
import time

import chromadb
import docx
from openai import OpenAI

# FIX #7: Shared helper — no longer duplicated here
from utils import clean_job_description, logger

if sys.stdout.encoding != "utf-8":
    sys.stdout.reconfigure(encoding="utf-8")

# ── Config ────────────────────────────────────────────────────────────────────
root_dir    = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
config_path = os.path.join(root_dir, "config.json")

if not os.path.exists(config_path):
    # FIX #1: raise instead of exit() — safe for import scenarios
    raise SystemExit("ERROR: config.json not found. Complete the Setup Wizard first.")

with open(config_path, "r", encoding="utf-8") as f:
    config_data = json.load(f)

api_key = config_data.get("groqApiKey") or os.getenv("GROQ_API_KEY", "")
if not api_key:
    raise SystemExit("ERROR: groqApiKey not found in config.json or GROQ_API_KEY env var.")

client = OpenAI(base_url="https://api.groq.com/openai/v1", api_key=api_key)

# ── Vector DB ─────────────────────────────────────────────────────────────────
logger.info("Connecting to Vector Database...")
db_path = os.path.join(root_dir, "backend", "chroma_db")
try:
    chroma_client = chromadb.PersistentClient(path=db_path)
    collection    = chroma_client.get_collection(name="career_projects")
except Exception as e:
    raise SystemExit(
        f"ERROR: Cannot connect to ChromaDB: {e}\n"
        "Run the Setup Wizard to embed your Career Profile first."
    )

# ── Jobs CSV ──────────────────────────────────────────────────────────────────
csv_file = os.path.join(root_dir, "jobs_database.csv")
if not os.path.isfile(csv_file):
    raise SystemExit("ERROR: jobs_database.csv not found. Run the Scraper first.")

output_dir = os.path.join(root_dir, "Tailored_Resumes")
os.makedirs(output_dir, exist_ok=True)

# FIX #2: Read master resume ONCE before the loop (was re-read every iteration)
master_resume_path = os.path.join(root_dir, "Master_Career_Profile.txt")
if not os.path.exists(master_resume_path):
    raise SystemExit("ERROR: Master_Career_Profile.txt not found.")

with open(master_resume_path, "r", encoding="utf-8") as f:
    base_resume = f.read()

logger.info("\n--- Starting End-to-End Resume Generation ---\n")

# ── Per-job constants ─────────────────────────────────────────────────────────
# FIX #5: Hard token limits to prevent silent context-window overflow on Groq free tier
MAX_JD_CHARS     = 3000
MAX_RESUME_CHARS = 4000

# FIX #3: Model list — fallback is automatic, loop does not need a continue hack
MODELS = ["llama-3.3-70b-versatile", "llama-3.1-8b-instant"]


def call_llm(messages: list, max_tokens: int = 1500,
             temperature: float = 0.5, json_mode: bool = False) -> str:
    """
    Try each model in MODELS in order. Raise RuntimeError if all fail.
    FIX #3: Clean model-fallback via list iteration (original continue logic was broken).
    """
    last_err = None
    for model in MODELS:
        try:
            kwargs: dict = dict(
                model=model,
                max_tokens=max_tokens,
                temperature=temperature,
                messages=messages,
            )
            if json_mode:
                kwargs["response_format"] = {"type": "json_object"}
            resp = client.chat.completions.create(**kwargs)
            return resp.choices[0].message.content
        except Exception as e:
            logger.warning("Model %s failed: %s. Trying next...", model, e)
            last_err = e
            time.sleep(5)  # brief pause before trying fallback

    raise RuntimeError(f"All models exhausted. Last error: {last_err}")


# ── Main processing loop ──────────────────────────────────────────────────────
with open(csv_file, mode="r", encoding="utf-8") as file:
    reader = csv.reader(file)
    header = next(reader, None)

    for row in reader:
        if len(row) < 5:
            continue

        job_title, company, link, apply_type = row[0], row[1], row[2], row[3]
        job_description = clean_job_description(row[4])

        # Only generate resumes for Company Website jobs
        if apply_type != "Company Website":
            continue

        safe_company = re.sub(r'[\\/*?:"<>|]', "_", company).replace(" ", "_")
        safe_title   = re.sub(r'[\\/*?:"<>|]', "_", job_title).replace(" ", "_")
        md_filename  = f"{safe_company}_{safe_title}.md"
        docx_filename = md_filename.replace(".md", ".docx")
        md_path      = os.path.join(output_dir, md_filename)
        docx_path    = os.path.join(output_dir, docx_filename)

        # Skip if already generated
        if os.path.exists(md_path) and os.path.exists(docx_path):
            logger.info("Skipping %s — resume already exists.", company)
            continue

        logger.info("Generating resume for %s — %s...", company, job_title)

        # ── RAG: retrieve top-3 relevant project chunks ───────────────────────
        logger.info("   -> Querying Vector DB...")
        try:
            results = collection.query(
                query_texts=[job_description[:MAX_JD_CHARS]], n_results=3
            )
            retrieved_projects = "\n\n".join(results["documents"][0])
        except Exception as rag_e:
            logger.warning("   -> RAG query failed: %s", rag_e)
            retrieved_projects = ""

        # FIX #4: Guard against empty RAG context — fall back to full resume
        if not retrieved_projects.strip():
            logger.warning(
                "   -> ⚠️  RAG returned no results. "
                "Falling back to full career profile."
            )
            retrieved_projects = base_resume[:MAX_RESUME_CHARS]

        # FIX #5: Truncate inputs before sending to LLM
        jd_truncated     = job_description[:MAX_JD_CHARS]
        resume_truncated = base_resume[:MAX_RESUME_CHARS]

        my_real_resume = (
            f"## CANDIDATE MASTER CAREER PROFILE\n{resume_truncated}\n\n"
            f"## HIGH-RELEVANCE FOCAL POINTS (from Vector DB for this JD)\n"
            f"{retrieved_projects}"
        )

        try:
            # ── Agent 1: Writer ───────────────────────────────────────────────
            writer_system = f"""You are an expert tech recruiter and master resume writer.
Craft a highly targeted 1-page resume using the candidate's Master Profile and focal points.

CRITICAL RULES:
1. SECTION HEADERS (use exactly these):
   {config_data.get("fullName", "Candidate Name")}
   {config_data.get("cityCountry", "City")} | {config_data.get("phone", "Phone")} | {config_data.get("email", "Email")} | {config_data.get("linkedin", "LinkedIn")}
   PROFESSIONAL SUMMARY
   CORE COMPETENCIES
   PROFESSIONAL EXPERIENCE
   INTERNSHIP
   EDUCATION
   CERTIFICATIONS & PROFESSIONAL DEVELOPMENT
   ADDITIONAL
2. RAG INTEGRATION: The HIGH-RELEVANCE FOCAL POINTS are real sections from the candidate's
   background matching this JD. Give them extra prominence and detail.
3. DYNAMIC ROLE TRANSLATION: Frame existing skills using {job_title} terminology.
4. STRICT HONESTY: Every skill/tool must exist in the Master Profile. No fabrication.
5. NO PLACEHOLDERS or extra notes at the bottom."""

            writer_user = (
                f"CANDIDATE PROFILE:\n{my_real_resume}\n\n"
                f"TARGET JOB DESCRIPTION:\n{jd_truncated}\n\n"
                "Write my tailored 1-page resume."
            )

            # FIX #3: Uses clean call_llm() with automatic model fallback
            tailored_content = call_llm(
                [{"role": "system", "content": writer_system},
                 {"role": "user",   "content": writer_user}],
                max_tokens=1500, temperature=0.5,
            )

            # ── Agent 2: ATS Scorer ───────────────────────────────────────────
            logger.info("   -> Scoring resume for %s...", company)

            scorer_system = """You are a cynical, critical ATS scanner and senior recruiter.
Evaluate the resume against the job description. Output strict JSON:
{
  "overall_score": <int 0-100>,
  "missing_keywords": [<list>],
  "feedback": "<2-3 actionable sentences>"
}

RULES:
1. KEYWORD MATCH: Keywords from JD must appear in the resume.
2. EVIDENCE CHECK: Deduct 10 pts per keyword listed without evidence in experience bullets.
3. DOMAIN ALIGNMENT: Only deduct heavily (<65) for highly specialised domains where
   the candidate has zero matching experience.
4. HONESTY: Scores above 85 require authentic, evidence-backed alignment."""

            scorer_user = (
                f"Job Description:\n{jd_truncated}\n\n"
                f"Candidate Resume:\n{tailored_content}"
            )

            raw_eval = call_llm(
                [{"role": "system", "content": scorer_system},
                 {"role": "user",   "content": scorer_user}],
                max_tokens=1000, temperature=0.0, json_mode=True,
            )

            # Parse ATS score safely
            try:
                json_match = re.search(r"\{.*\}", raw_eval, re.DOTALL)
                eval_data  = json.loads(json_match.group() if json_match else raw_eval)
                raw_score  = eval_data.get("overall_score", 0)
                # FIX #6: Normalise score to int with safe fallback
                score    = int(raw_score) if str(raw_score).strip().isdigit() else 0
                missing  = ", ".join(eval_data.get("missing_keywords", []))
                feedback = eval_data.get("feedback", "No suggestions provided.")
            except (json.JSONDecodeError, ValueError) as json_err:
                logger.warning("   -> ⚠️  Failed to parse ATS JSON: %s", json_err)
                score, missing, feedback = 0, "Unknown", "LLM failed to provide structured feedback."

            # ── Save Markdown ─────────────────────────────────────────────────
            with open(md_path, "w", encoding="utf-8") as out:
                out.write(f"**Target Company:** {company}\n")
                out.write(f"**Job Link:** {link}\n")
                out.write(f"**Apply Type:** {apply_type}\n")
                out.write(f"**ATS Score:** {score}/100\n")
                out.write(f"**Missing Keywords:** {missing}\n")
                out.write(f"**Suggestions:** {feedback}\n\n---\n\n")
                out.write(tailored_content)

            # ── Save DOCX ─────────────────────────────────────────────────────
            doc = docx.Document()
            title_p = doc.add_paragraph(f"Resume: {company} — {job_title}")
            if title_p.runs:
                title_p.runs[0].bold = True

            def add_md_paragraph(doc_obj, text, style=None):
                p = doc_obj.add_paragraph(style=style)
                for part in re.split(r"(\*\*.*?\*\*)", text):
                    if part.startswith("**") and part.endswith("**"):
                        p.add_run(part[2:-2]).bold = True
                    else:
                        p.add_run(part)
                return p

            for line in tailored_content.split("\n"):
                stripped = line.strip()
                if stripped.startswith("# "):
                    p = add_md_paragraph(doc, stripped[2:])
                    for r in p.runs:
                        r.bold = True
                elif stripped.startswith("## "):
                    p = add_md_paragraph(doc, stripped[3:])
                    for r in p.runs:
                        r.bold = True
                elif stripped.startswith("### "):
                    p = add_md_paragraph(doc, stripped[4:])
                    for r in p.runs:
                        r.bold = True
                elif stripped.startswith(("- ", "* ")):
                    try:
                        add_md_paragraph(doc, stripped[2:], style="List Paragraph")
                    except Exception:
                        try:
                            add_md_paragraph(doc, stripped[2:], style="List Bullet")
                        except Exception:
                            add_md_paragraph(doc, "• " + stripped[2:])
                elif stripped:
                    add_md_paragraph(doc, stripped)

            doc.save(docx_path)
            logger.info(
                "-> ✅ SUCCESS! Score: %d/100. Saved: Tailored_Resumes/%s",
                score, docx_filename,
            )

            # Anti-rate-limit pause (fine as synchronous — subprocess context)
            time.sleep(3)

        except RuntimeError as e:
            logger.error("-> ❌ All LLM models failed for %s: %s", company, e)
        except Exception as e:
            logger.error("-> ❌ Unexpected error for %s: %s", company, e)

logger.info("All done! Open the Tailored_Resumes folder to see your files.")
